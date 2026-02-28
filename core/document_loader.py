import os
from core.utils.security_guard import security_guard

class DocumentLoader:
    """
    A simple loader to read text from files in the KnowledgeBase.
    Requires: pip install pypdf
    """
    def __init__(self, knowledge_base_path="01_Teaching_Modules/KnowledgeBase"):
        self.kb_path = knowledge_base_path
    
    def load_documents(self):
        documents = []
        if not os.path.exists(self.kb_path):
            os.makedirs(self.kb_path)
            return documents
        
        # Enforce security scoping
        for filename in security_guard.secure_list_dir(self.kb_path):
            file_path = os.path.join(self.kb_path, filename)
            content = ""
            
            try:
                if filename.endswith((".txt", ".md")):
                    with security_guard.secure_open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif filename.endswith(".pdf"):
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                elif filename.endswith(".docx"):
                    import docx
                    doc = docx.Document(file_path)
                    content = "\n".join([p.text for p in doc.paragraphs])
                    for table in doc.tables:
                        for row in table.rows:
                            content += "\n" + "\t".join([cell.text.replace('\n', ' ') for cell in row.cells])
                
                # Try high-fidelity parsing with Docling if available
                if filename.endswith((".pdf", ".docx")):
                    try:
                        from docling.document_converter import DocumentConverter
                        print(f"  [DocumentLoader] Using Docling for high-fidelity parsing of {filename}...")
                        converter = DocumentConverter()
                        result = converter.convert(file_path)
                        docling_content = result.document.export_to_markdown()
                        if docling_content:
                            content = docling_content
                    except ImportError:
                        pass # Fallback to standard pypdf/docx already handled above
                    except Exception as e:
                        print(f"  [DocumentLoader] Docling error for {filename}: {e}")
                
                # Check if content extraction yielded nothing (common for scanned PDFs)
                if not content.strip() and filename.endswith(".pdf"):
                    print(f"  [DocumentLoader] PDF content empty for {filename}. Trying Multimodal extraction...")
                    try:
                        from core.utils.multimodal_extractor import MultimodalExtractor
                        extractor = MultimodalExtractor()
                        content = extractor.process_file(file_path)
                    except Exception as me:
                        print(f"  [DocumentLoader] Multimodal extraction failed: {me}")

                if content:
                    documents.append({
                        "filename": filename,
                        "content": content
                    })
                    
            except Exception as e:
                print(f"Error reading {filename}: {e}")
                
        return documents

def process_document(source):
    """Standalone function requested by night_shift.py"""
    loader = DocumentLoader(knowledge_base_path=os.path.dirname(source) if source else "01_Teaching_Modules/KnowledgeBase")
    if source and os.path.exists(source):
        # Specific file handling
        content = ""
        filename = os.path.basename(source)
        try:
            if filename.endswith((".txt", ".md")):
                with open(source, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif filename.endswith(".pdf"):
                from pypdf import PdfReader
                reader = PdfReader(source)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            elif filename.endswith(".docx"):
                import docx
                doc = docx.Document(source)
                content = "\n".join([p.text for p in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        content += "\n" + "\t".join([cell.text.replace('\n', ' ') for cell in row.cells])
            
            # High-fidelity fallback/upgrade
            if filename.endswith((".pdf", ".docx")):
                try:
                    from docling.document_converter import DocumentConverter
                    print(f"  [process_document] Upgrading to Docling for {filename}...")
                    converter = DocumentConverter()
                    result = converter.convert(source)
                    docling_content = result.document.export_to_markdown()
                    if docling_content:
                        content = docling_content
                except (ImportError, Exception):
                    # Keep previous extraction text if docling fails or isn't installed
                    pass
            
            # Fallback for scanned PDFs
            if not content.strip() and filename.endswith(".pdf"):
                print(f"  [process_document] PDF empty for {filename}. Using Multimodal fallback...")
                from core.utils.multimodal_extractor import MultimodalExtractor
                extractor = MultimodalExtractor()
                content = extractor.process_file(source)

            return content
        except Exception as e:
            return f"Error processing version standalone: {e}"
    
    # If no specific source or directory-wide fallback
    docs = loader.load_documents()
    return "\n\n".join([d['content'] for d in docs])
