import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_paper(output_path):
    doc = docx.Document()
    
    # --- Style Settings ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # --- Header ---
    title = doc.add_heading('Hybrid Multi-Criteria Decision-Making Model for Retail Credit Risk Assessment: Evidence from Vietnam', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    authors = doc.add_paragraph('Nguyen Nam Trung, Nguyen Thi Hong Thuy, Luu Quoc Dat, Dang Thu Hang, Bui Phuong Chi')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    aff = doc.add_paragraph('VNU University of Economics and Business, Hanoi 100000, Vietnam')
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "Purpose: This study develops and validates a hybrid multi-criteria decision-making (MCDM) model combining Fuzzy Analytic Hierarchy Process (Fuzzy-AHP) and Technique for Order Preference by Similarity to Ideal Solution (TOPSIS) for individual credit risk assessment in data-scarce environments. "
        "Design/Methodology/Approach: We employ a two-stage mixed-methods approach. First, expert elicitation via Fuzzy-AHP establishes criterion weights across four risk dimensions (Credit History, Financial Capacity, Personal Stability, and Loan Profile). Second, TOPSIS scores 340 consumer loan applications from Vietnamese commercial banks (2022-2024). Model performance is benchmarked against Logistic Regression using AUC-ROC, precision, recall, and F1-score metrics. "
        "Findings: The Fuzzy-AHP-TOPSIS model demonstrates superior discriminatory power (AUC=0.96) compared to Logistic Regression (AUC=0.93). Critically, the proposed model achieves 93% recall versus 79% for the benchmark, reducing Type II errors (missed defaults) by 17.7%. CIC credit score emerges as the dominant criterion (42.7% weight). "
        "Practical Implications: The transparent 'white-box' architecture enables explainable AI (XAI) compliance with Vietnam's Decree 13/2023/ND-CP. We propose a tiered risk assessment framework: automated scoring for standard applications and MCDM for thin-file or 'Gray-Zone' borrowers. "
        "Originality/Value: This is the first empirical validation of hybrid MCDM for retail credit scoring in Vietnam, addressing the critical gap between advanced machine learning opacity and regulatory transparency requirements."
    )
    doc.add_paragraph('Keywords: credit risk assessment, multi-criteria decision making, Fuzzy-AHP, TOPSIS, thin-file borrowers, explainable AI, emerging Markets')
    
    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=1)
    p1 = doc.add_paragraph()
    p1.add_run(
        "Emerging markets like Vietnam face significant credit rationing due to acute information asymmetry. "
        "This creates a 'Market for Lemons' (Akerlof, 1970), where lenders cannot distinguish between solvent and insolvent 'thin-file' borrowers. "
        "While Machine Learning (ML) offers predictive power, the 'black-box' nature of these models often violates emerging transparency regulations like Vietnam's Decree 13/2023/ND-CP. "
        "There is a critical need for a 'white-box' approach that integrates expert knowledge with empirical data to ensure both accuracy and accountability."
    )

    # --- 3. Methodology ---
    doc.add_heading('3. Methodology', level=1)
    
    doc.add_paragraph(
        "Figure 1. Research methodology framework. The framework illustrates a two-stage process: (1) Knowledge-based weighting via Fuzzy AHP to capture expert risk perception, and (2) Data-driven scoring via TOPSIS where loan applications are ranked by their geometric distance to an ideal risk profile."
    )

    doc.add_heading('3.2. Data Preprocessing and Encoding', level=2)
    doc.add_paragraph(
        "A critical refinement in our methodology is the use of the 'Benefit Principle' for categorical encoding in the MCDM model. "
        "Unlike purely statistical models that use one-hot encoding (treating all categories as neutral), we utilize Ordinal Encoding derived from expert consensus. "
        "For example, 'Homeowner' (3) is assigned a higher value than 'Renter' (1) because it signals a financial safety net, reducing the probability of flight. "
        "This knowledge-driven encoding stabilizes the model even when specific categories are under-represented in the small sample (N=340), preventing the statistical underfitting common in purely data-driven approaches."
    )

    doc.add_paragraph(
        "Figure 2. Data preprocessing workflow. This flowchart illustrates the parallel paths: Path A for the knowledge-driven MCDM model and Path B for the statistical Logistic Regression baseline. It highlights the divergence in feature selection: Path A relies on expert pillars, while Path B uses a recursive backward elimination funnel based on AIC."
    )

    doc.add_paragraph(
        "Figure 3. Encoding strategies for categorical variables. The diagram contrasts how variables like 'Job Title' and 'Residence' are handled: linear binary encoding for Logistic Regression to preserve data patterns versus hierarchical ordinal encoding for MCDM to reflect expert-defined risk logic."
    )

    # Table 2
    doc.add_heading('Table 2. Fuzzy AHP Global Weights', level=2)
    t2 = doc.add_table(rows=6, cols=3)
    t2.style = 'Table Grid'
    hdr = t2.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Pillar', 'Criterion', 'Crisp Weight'
    data2 = [
        ('Credit History', 'CIC Score (D01)', '0.427'),
        ('Financial Capacity', 'DTI (D04), Monthly Income (C06)', '0.285'),
        ('Personal Stability', 'Job Tenure (C03), Residential Status (B06)', '0.180'),
        ('Loan Profile', 'Loan Amount (A05), Loan Term (A06)', '0.108'),
        ('Total', '', '1.000')
    ]
    for i, row_data in enumerate(data2):
        row = t2.rows[i+1].cells
        row[0].text, row[1].text, row[2].text = row_data

    # --- 5. Discussion ---
    doc.add_heading('5. Discussion', level=1)
    
    doc.add_heading('5.1. Knowledge Compensation and Safety Bias', level=2)
    doc.add_paragraph(
        "The empirical results validate the efficiency of knowledge-driven approaches in data-constrained environments. "
        "The MCDM model's superior Recall (93% vs 79%) indicates a strategic 'Safety Bias'. "
        "In the context of the 'Market for Lemons' (Akerlof, 1970), the higher cost of Type II errors (missed defaults) compared to Type I errors (rejected good loans) justifies this conservative approach. "
        "By integrating the 42.7% weight for CIC scores, the model ensures that historical delinquency is penalized more heavily than purely statistical coefficients might suggest in a sparse dataset."
    )

    doc.add_heading('5.2. Explainable AI and Regulatory Compliance', level=2)
    doc.add_paragraph(
        "Beyond predictive accuracy, the proposed framework addresses the 'Black Box' dilemma. "
        "With the enforcement of Vietnam’s Decree 13/2023/ND-CP on personal data protection, financial institutions are increasingly required to provide explainable logic for automated decisions. "
        "The geometric nature of TOPSIS offers 'Structural Transparency': a loan officer can explicitly explain to an applicant that their score was penalised due to the Euclidean distance in the 'Personal Stability' criterion. "
        "This transparency bridges the gap between high-accuracy automated scoring and the 'right to explanation' mandated by law."
    )

    doc.add_paragraph(
        "Figure 4. ROC Curve Comparison. The visualization shows the dominant performance of the Fuzzy AHP-TOPSIS model (AUC=0.96) over the Logistic Regression benchmark (AUC=0.93), particularly in the low False Positive Rate region."
    )

    doc.add_paragraph(
        "Figure 5. Tiered Risk Assessment Framework. We propose using automated statistical scoring for standard applications (Tier 1) and routing 'Gray-Zone' or thin-file applications to the more transparent, expert-weighted MCDM model (Tier 2) to ensure safety and explainability."
    )

    # --- References ---
    doc.add_heading('6. References', level=1)
    refs = [
        "Akerlof, G. A. (1970). The Market for 'Lemons': Quality Uncertainty and the Market Mechanism. The Quarterly Journal of Economics, 84(3), 488–500.",
        "Adadi, A., & Berrada, M. (2018). Peeking Inside the Black-Box: A Survey on Explainable Artificial Intelligence (XAI). IEEE Access, 6, 52138–52160.",
        "Chang, D.-Y. (1996). Applications of the extent analysis method on fuzzy AHP. European Journal of Operational Research, 95(3), 649–655.",
        "Government of Vietnam. (2023). Decree No. 13/2023/ND-CP on Personal Data Protection.",
        "State Bank of Vietnam. (2021). Circular No. 11/2021/TT-NHNN on classification of assets and risk provisions.",
        "Thomas, L. C., Edelman, D. B., & Crook, J. N. (2002). Credit Scoring and Its Applications. SIAM.",
        "World Bank. (2022). The Global Findex Database 2021: Financial Inclusion, Digital Payments, and Resilience."
    ]
    for r in refs:
        doc.add_paragraph(r)
    
    # --- Appendixes ---
    doc.add_heading('Appendixes', level=1)
    doc.add_heading('Appendix A: Sample pairwise comparison questionnaire', level=2)
    doc.add_paragraph("Linguistic terms were used: Equally Important (1), Moderately More (3), Strongly More (5), Very Strongly More (7), Extremely More (9).")
    doc.add_heading('Appendix B: Calculation walkthrough', level=2)
    doc.add_paragraph("For an applicant with CIC=650, Income=20M, DTI=0.3: (1) Normalize values, (2) Apply weights (0.427 for CIC...), (3) Calculate Euclidean distance to Ideal (1,1,1) and Worst (0,0,0) profiles.")

    doc.save(output_path)
    print(f"Deeply refined paper saved to {output_path}")

if __name__ == "__main__":
    output_path = r"C:\Users\Trung Nguyen\Desktop\Working data\Reseach Paper - MCDM for retail credit rating in Vietnam - Final Polished.docx"
    create_final_paper(output_path)
