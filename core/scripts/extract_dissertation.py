import docx
import os

def extract_text(docx_path, output_path):
    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} not found.")
        return
    
    doc = docx.Document(docx_path)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(full_text))
    print(f"Successfully extracted text to {output_path}")

if __name__ == "__main__":
    docx_path = r"C:\Users\Trung Nguyen\Desktop\Hồ sơ bảo vệ cơ sở\Luận án-draft-v3-citation.docx"
    output_path = r"C:\Users\Trung Nguyen\Desktop\Hồ sơ bảo vệ cơ sở\dissertation_content.txt"
    extract_text(docx_path, output_path)
